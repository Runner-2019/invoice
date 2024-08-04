import os
import wx
import re
import imghdr
from pubsub      import pub
from datetime    import datetime
from paddleocr   import PaddleOCR
from docx        import Document
from docx.shared import Inches

def is_image_file(filename) -> bool:
    img_type = imghdr.what(filename)
    return img_type is not None

def count_files_in_directory(directory_path) -> int:
    return len([entry for entry in os.listdir(directory_path) if is_image_file(os.path.join(directory_path, entry))])

# need to run only once to download and load model into memory
OCR = PaddleOCR( show_log=False, use_angle_cls=True, lang="ch")

class OCRHandle:
    def __init__(self, invoices_path: str, prompt: list):
        self.invoices_path = invoices_path
        self.prompt = prompt
        self.total_invoices_count = 0
        self.total_error_invoice_list=[]
        self.total_account_list = []      # invoice account of all invoice pictures
        self.cur_invoice_path = ""        # invoice picture which is currently handling
        self.cur_invoice_index = 0        # invoice picture index of all invoices
        self.cur_account = 0              # invoice account
        self.cur_account_str = ""         # invoice account string
        self.cur_text_list = []           # invoice picture to text
        self.log_file = ""                # detailed log in docx format
        self.document = None              # python docx
        self.succ_para = None             # success paragraph
        self.fail_para = None
        self.conclude_para  = None


    # Convert picture to text list, discards box coordinates and predicate score.
    def __convert_once(self):
        result = OCR.ocr(self.cur_invoice_path, cls=True)
        # BUG: why result[0]?
        for index, line in enumerate(result[0]):
            text=line[1][0]
            self.cur_text_list.append(text)


    # Check whether text lists  is valid
    def __check_text_list(self):
        if not self.cur_text_list:
            raise RuntimeError("without text list")


    # Extract account number from text list
    def __get_cur_account(self):
        account_str_list = []
        for text in self.cur_text_list:
            for prompt in self.prompt:
                if text.find(prompt) != -1:
                    account_str_list.append(text)

        if len(account_str_list) == 0:
            raise RuntimeError("without account string")

        if len(account_str_list) > 1:
            raise RuntimeError("too many account strings")

        # exactly one
        self.cur_account_str = account_str_list[0]
        print(f'Index: {self.cur_invoice_index}, account string: {self.cur_account_str} of invoice picture: {self.cur_invoice_path}')

        numbers = re.findall(r'\d+\.\d+|\d+', self.cur_account_str)
        if len(numbers) != 1:
            raise RuntimeError("too many accounts from account string")

        self.cur_account = float(numbers[0])
        print(f'Index: {self.cur_invoice_index}, account number: {self.cur_account} of invoice picture: {self.cur_invoice_path}')


    def __creat_log_file(self):
        path="./results"
        if not os.path.exists(path):
            os.makedirs(path)

        current_datetime_str = datetime.now().strftime("%Y_%m_%d_%H_%M_%S")
        filename = current_datetime_str
        while os.path.exists(os.path.join(path, filename)):
            counter += 1
            filename = f"{base_filename}_no_{counter}"

        file_path = f"{path}/{filename}.doc"
        with open(file_path, 'w') as file:
            pass # create

        self.log_file  = file_path
        self.document  = Document()
        self.succ_para = self.document.add_heading("1.正常解析的发票\n")
        self.fail_para = self.document.add_heading("2.解析失败的发票\n")
        self.conclude_para = self.document.add_heading("3.总结\n")


    # if is_success, then append current picture to success part.
    def __append_to_detail_log(self, is_success):
        if not self.log_file:
            self.__creat_log_file()

        if is_success:
            # self.document.add_picture(self.cur_invoice_path, width=Inches(6))
            # self.document.add_paragraph(f"result: {self.cur_account}")
            # self.document.add_paragraph("\n")

            self.succ_para.add_run().add_picture(self.cur_invoice_path, width=Inches(6))
            self.succ_para.add_run(f"上图中的发票金额: {self.cur_account}")
            self.succ_para.add_run("\n")

    def __make_conclue(self):
        formula = ""
        for index, account in enumerate(self.total_account_list):
            if index == 0:
                formula += f"{account}"
            else:
                formula += f" + {account}"
        formula += f" = {sum(self.total_account_list)}"
        self.conclude_para.add_run("正常解析的发票的计算公式:\n")
        self.conclude_para.add_run(formula)
        self.conclude_para.add_run("\n")

        if not self.total_error_invoice_list:
            self.fail_para.add_run("无\n")
        else:
            self.conclude_para.add_run("下面的文件解析失败了:\n")
            for invoice in self.total_error_invoice_list:
                self.fail_para.add_run().add_picture(invoice, width=Inches(6))
                self.fail_para.add_run("\n")


    def __finish_one_picture(self):
        print("__finish_one_picture")
        self.total_account_list.append(self.cur_account)
        self.cur_account = 0
        self.cur_invoice_path = ""
        self.cur_account_str = ""
        self.cur_text_list = []

    def run(self) -> None:
        if not os.path.isdir(self.invoices_path):
            raise RuntimeError(f"提供的发票文件夹有误: {self.invoices_path}")

        self.total_invoices_count = count_files_in_directory(self.invoices_path)
        print(f"total invoces count: {self.total_invoices_count}")
        if self.total_invoices_count == 0:
            raise RuntimeError(f"提供的发票文件夹中不包含发票图片")

        for filename in os.listdir(self.invoices_path):
            invoice = os.path.join(self.invoices_path, filename)
            if os.path.isfile(invoice):
                if not is_image_file(invoice):
                    continue
                try:
                    print(f"dealing: {invoice}")
                    self.cur_invoice_path = invoice
                    self.cur_invoice_index += 1
                    self.__convert_once()
                    self.__check_text_list()
                    self.__get_cur_account()
                    self.__append_to_detail_log(True)
                    cur_progress = int(self.cur_invoice_index * 1.0 / self.total_invoices_count * 100)
                    wx.CallAfter(pub.sendMessage, "update_process", count=cur_progress)
                except Exception as e:
                    print(str(e))
                    self.total_error_invoice_list.append(self.cur_invoice_path)
                    self.__append_to_detail_log(False)
                self.__finish_one_picture()

        print(f'start to write to file.')
        self.__make_conclue()
        self.document.save(self.log_file)

        total_account = sum(self.total_account_list)
        print(f'total account of all tickets: {total_account}')
        wx.CallAfter(pub.sendMessage,
                     "finish_compute",
                     result=total_account,
                     detailed_result_file=self.log_file)
